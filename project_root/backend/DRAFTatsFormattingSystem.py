# See markdown documents for overview and todolsit

import spacy
from spacy_wordnet.wordnet_annotator import WordnetAnnotator
from collections import Counter
from docx import Document
import requests
from bs4 import BeautifulSoup
import os
import logging
import subprocess
import sys
from sklearn.feature_extraction.text import TfidfVectorizer 
#TODO: import pandas as pd #SWITCH TO GOOGLESHEET OR EXCEL IDK, make gsheet first and backend then transfer to excel

# Global loading for efficiency, If I reload it every time I call a function, I'll waste resources.
logging.basicConfig(level=logging.INFO)
nlp = spacy.load("en_core_web_md")

#Default Settings
#TODO use configiration file for defailt settings
resume_doc_filepath: str = "resume.docx"
job_post_url: str = None
default_testing: bool = True
    
# BASIC TESTING STUFF
basic_job_description = """We are seeking a highly motivated Data Analyst to join our team. The ideal candidate will have 
a strong background in data analytics, statistical modeling, and visualization. You will be 
responsible for analyzing complex datasets, generating actionable insights, and supporting 
strategic decision-making processes.

Key Responsibilities:
- Analyze large datasets to identify trends and patterns.
- Create interactive dashboards using tools like Tableau or Power BI.
- Collaborate with cross-functional teams to understand business needs.
- Develop predictive models using Python or R.
- Write SQL queries to extract and transform data from relational databases.

Qualifications:
- Bachelor's degree in Statistics, Mathematics, Computer Science, or a related field.
- Proficiency in Python, R, or other statistical programming languages.
- Experience with SQL for data extraction and manipulation.
- Familiarity with Tableau, Power BI, or similar visualization tools.
- Strong communication and presentation skills."""

basic_resume_text = """John Doe
Data Analyst
johndoe@example.com | LinkedIn: linkedin.com/in/johndoe | Phone: (123) 456-7890

Professional Summary:
Detail-oriented data analyst with 3+ years of experience in interpreting and analyzing data 
to drive successful business solutions. Adept at building dashboards, writing complex SQL 
queries, and developing predictive models to solve challenging problems.

Skills:
- Data Visualization: Tableau, Power BI, matplotlib, seaborn
- Statistical Analysis: Python, R, Excel
- Database Management: SQL, MySQL, PostgreSQL
- Machine Learning: Scikit-learn, TensorFlow
- Business Intelligence: Building dashboards, reports, and ad hoc analysis

Work Experience:
Data Analyst | XYZ Corporation | Jan 2020 - Present
- Designed and maintained interactive dashboards for tracking key business metrics.
- Conducted statistical analysis to support marketing strategies, resulting in a 15% increase in ROI.
- Developed machine learning models to predict customer churn, achieving an 85% accuracy rate.
- Collaborated with IT and business teams to streamline data pipelines.

Junior Data Analyst | ABC Tech | Jun 2018 - Dec 2019
- Analyzed customer behavior using SQL and Python to optimize product offerings.
- Automated data extraction processes, reducing report generation time by 30%.
- Created visualizations to communicate findings to stakeholders.

Education:
- Bachelor of Science in Statistics | University of Data Science | 2018

Certifications:
- Tableau Desktop Specialist
- Google Data Analytics Professional Certificate"""


#BASIC DATA #eventually maybe make keyword similarity so can put skills by job and find the job / skills if smiliar liek a giant database / tree IDK getting to complicated
predefined_keywords = {
    "tools": ["Python", "SQL", "Tableau", "Power BI"],
    "skills": ["data visualization", "statistical analysis", "predictive modeling"],
    "certifications": ["Google Data Analytics", "Tableau Certified Professional"],
    "qualifications": ["Bachelor's degree", "Statistics", "Computer Science"]
}

high_relevance_keywords = {
    "python", "sql", "tableau", "power bi", "data visualization",
    "predictive modeling", "machine learning", "statistics", "dashboard",
    "database", "excel", "scikit-learn", "r", "etl", "data pipeline"
}

"""Welcome to the Resume Optimizer!
This program helps you:
1. Parse your resume and a job description.
2. Identify keywords to improve your match score.
3. Generate a tailored LLM (ChatGPT) prompt to rewrite your resume.
4. Save and manage your updated resume automatically.

You can:
- Enter your resume and job description directly.
- Use default testing mode for examples.
- Save results to a .docx file.

Let's get started!"""

# INPUTS
def extract_text_from_docx(file_path):
    """
    Extracts and returns the text from a Word document.
    """
    
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

#TODO Develop
def extract_text_from_job_description(url):
    """
    Extracts and returns the text from a job posting URL.
    """
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Example: Extract text from a specific HTML tag
        # Modify the tag and class based on the structure of the job website
        job_description_text = soup.find_all(['p', 'div'], class_=None)
        
        # Join all extracted text into a single string
        text = " ".join([element.get_text() for element in job_description_text])
        return text
    else:
        return f"Failed to fetch URL. Status code: {response.status_code}"

# BACKEND / INTERNAL PROCESSES # Why: Seperation of Concerns / Single Responsibility, Maintainability, Testability, Reusability



def setup_program(intro_text):
    """
    Display introductory information about the program.
    
Current Pain Point: Users must paste their data multiple times, which can be streamlined by:
Automatically saving to a file after generating outputs.
Providing clearer instructions upfront.
Reducing repeated prompts.

Use a setup function to explain the program.
Allow a single input flow with an optional override."""

#add fucntionality here


def extract_keywords(text):
    """
    Extracts keywords (nouns, verbs, and adjectives) from the input text.
    """
    doc = nlp(text)
    keywords = [token.text.lower() for token in doc if token.pos_ in ("NOUN", "VERB", "ADJ")]
    return Counter(keywords)

def extract_keywords_table(doc, verbose=False) -> str:
        """
        Extracts keywords and optionally generates a table of extracted keywords.
        
        Args:
        - doc (spacy.tokens.Doc): A SpaCy Doc object.

        Returns:
        - output (str): A formatted string representing the keyword table.
        """
        # Initialize the table with headers
        output = """TESTING - Extracted Keywords Table:\n"""

        # Loop through tokens in the Doc and append rows to the table
        for token in doc:
            output += f"TEXT: {token.text} | pos: {token.pos_}\n"
    
        #print("Filtered Kywords test", list(filter(lambda x: x.pos_ in ("NOUN", "VERB", "ADJ"), doc)))
        return output

    # Assuming `doc` is already a processed SpaCy Doc object
    #print(extract_keywords_table(doc))
    
def get_synonyms(word):
    """
    Retrieves synonyms for a word using SpaCy's WordNet extension.
    Args:
        word (str): The word to find synonyms for.
    Returns:
        set: A set of synonyms.
    """
    token = nlp(word)[0]  # Process the word
    return {lemma for synset in token._.wordnet.synsets() for lemma in synset.lemma_names()}


def get_high_relevance_keywords(job_description, all_documents):
    """Feed job descriptions into the function and identify terms that are both frequent and unique.
    Extract words from the job description and compare them to a reference corpus to identify unique and relevant terms.
    TODO GET A CORPUS TOGETHER, this needs to be another program that everyonce in a while or by use guidnace goes and selects a bucnh of revlenat pages and find these words and redoes or grows the data everyonce and a while to keep up to date
    
    #FREQ BASED
    Extract Words:
    Tokenize the job description and other related job postings.
    Calculate Importance:
    Words that appear frequently in the job description but not in generic contexts (e.g., stopwords or filler words) are considered more relevant.
    """
    vectorizer = TfidfVectorizer(stop_words="english")
    tfidf_matrix = vectorizer.fit_transform(all_documents)
    feature_names = vectorizer.get_feature_names_out()
    
    # Focus only on the current job description
    job_tfidf = tfidf_matrix[0].toarray()[0]
    
    # Rank words by TF-IDF scores
    keywords = [(feature_names[i], job_tfidf[i]) for i in range(len(feature_names))]
    keywords = sorted(keywords, key=lambda x: -x[1])  # Sort by relevance
    return [word for word, score in keywords if score > 0.1]  # Filter by threshold

def get_relevant_keywords(job_description, high_value_terms):
    """Semantic Similarity:
    Identify terms in the job description or resume that are semantically similar to high-priority keywords.
    Compare each word in a job description to predefined high-priority terms (e.g., “Python,” “SQL”).
    
    Define high-priority words (e.g., "Python," "SQL").
    Compute the similarity of each word in the job description and resume to these terms.
    Retain keywords with high similarity scores.
    
    #RELEVANCE BASED
    """

    
    doc = nlp(job_description)
    relevant_keywords = set()
    
    for token in doc:
        for term in high_value_terms:
            similarity = nlp(token.text).similarity(nlp(term))
            if similarity > 0.75:  # Adjust threshold as needed
                relevant_keywords.add(token.text)
    return relevant_keywords

def extract_common_terms(job_descriptions):
    all_words = []
    for description in job_descriptions:
        words = description.lower().split()
        all_words.extend(words)
    
    # Filter out common stopwords
    stopwords = {"and", "or", "the", "to", "a", "of", "in"}
    filtered_words = [word for word in all_words if word not in stopwords]
    
    return Counter(filtered_words).most_common(10)  # Top 10 terms

def filter_relevant_keywords(keywords, relevance_criteria):
    """
    Filters keywords based on relevance criteria.
    
    Args:
    - keywords (set): The set of keywords to filter.
    - relevance_criteria (list): A list of high-relevance keyword categories (e.g., skills, tools).
    
    Returns:
    - filtered_keywords (set): A set of keywords that match the relevance criteria.
    """
    # Example criteria: Modify as needed
    high_relevance_keywords = {"tools", "technologies", "skills", "data", "analysis", "visualization",
                               "python", "sql", "tableau", "power bi", "predictive", "machine learning"}
    
    # Filter the keywords
    filtered_keywords = {word for word in keywords if word.lower() in high_relevance_keywords}
    return filtered_keywords

#TODO finish and make better
def compare_keywords(job_keywords, resume_keywords):
    """
    Compares job description keywords with resume keywords.
    Returns a match score and missing keywords.
    """
    job_set = set(job_keywords.keys())
    resume_set = set(resume_keywords.keys())
    
    #MODIFY TO CATCH / CONSIDER FOR SYNONMS?? but not semantics cus jsut too broad
    matched_keywords = job_set & resume_set
    missing_keywords = job_set - resume_set
    
    # Filter matched and missing keywords
    relevant_matched = filter_relevant_keywords(matched_keywords, high_relevance_keywords)
    relevant_missing = filter_relevant_keywords(missing_keywords, high_relevance_keywords)

    match_score = len(matched_keywords) / len(job_set) * 100
    return match_score, matched_keywords, missing_keywords

#TODO maybe eventuall be able to have go thru shcool recommended program linkend in , handshake jobs websites, and own internal data base and thin and find of new companeisn to go and target auto
def scrape_job_postings(url):
    """
    Scrape job postings from a given URL.
    """
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        # Adjust these selectors based on the website structure
        jobs = soup.find_all('div', class_='job-title')
        job_descriptions = [job.get_text() for job in jobs]
        return job_descriptions
    else:
        print("Failed to fetch job postings.")
        return []

def categorize_keywords(keywords, high_relevance):
    """ Provide actionable feedback by categorizing matched and missing keywords into high and low relevance.
    High Relevance: Skills, tools, certifications, and qualifications. 
    Low Relevance: Generic terms, structural words, or filler phrases."""
    high = {word for word in keywords if word in high_relevance}
    low = keywords - high
    return high, low


# Empowers users to use tools like ChatGPT, without me having integrating an API which can be costly or have the program run an oncomputer LLM which is unrealistic and time consuming 
def generate_llm_prompt(resume_text, missing_keywords):
    """
    Generates a prompt for a user to optimize their aresume with only relevant keywords using an LLM like ChatGPT.

    Args:
    - resume_text (str): The current text of the user's resume.
    - missing_keywords (set): A set of missing keywords.

    Returns:
    - str: A formatted LLM prompt.
    """
    prompt = f"""
I have the following resume:

{resume_text}

However, it is missing the following important keywords relevant to a job description I am targeting:

{', '.join(missing_keywords)}

Please rewrite my resume to include these keywords in a natural, professional, and concise way while maintaining readability. Ensure the updated resume still accurately reflects my skills and experience.
"""
    return prompt

def output_results(resume_text, job_description_text, match_score, matched_keywords, missing_keywords, resume_doc_filepath=None, job_post_url=None) -> str:
    """outputs results and feedback from analysis"""
    results = f"""
PROGRAM OUTPUT:
{'─' * 80}

INPUTS:

Resume:
Resume File Path: {resume_doc_filepath}
{resume_text}

Job Description:
Job Post URL: {job_post_url}
{job_description_text}

{'─' * 80}

ANALYSIS:

RESULTS:
Match Score: {match_score:.2f}%
Matched Keywords:
{', '.join(matched_keywords)}
Missing Keywords:
{', '.join(missing_keywords)}

{'─' * 80}

Recommended LLM / Chat GPT prompt to rewrite your resume to improve your score
LLM PROMPT:
{generate_llm_prompt(resume_text, missing_keywords)}

{'─' * 80}
    """
    
    return results

# WHY -  Helps reduce user painpoint of opening, closing, editing, copy and pasting tons of documetns - makes more convenient. Allow users to paste directly into the program and update their resume dynamically.
def update_resume_file(existing_doc, gpt_output):
    """
    Append GPT output to an existing resume file.
    """
    doc = Document(existing_doc)
    doc.add_paragraph("\nGPT Suggested Updates:")
    doc.add_paragraph(gpt_output)
    updated_file = "Updated_" + existing_doc
    doc.save(updated_file)
    print(f"Resume updated and saved to {updated_file}")

# WHY - For convenience and to stay organized, mainly for me. Need to have option to specify ghseets or excel and use api or whatever. Also have settings menu and tutorial but that will come after funcionality
#Create an Excel-based tracker for job applications.

def update_application_tracker():
    """
    Create or update a job application tracker.
    """
    tracker_file = "Application_Tracker.xlsx"

    # Example data
    data = {
        "Job Title": ["Data Analyst", "Python Developer"],
        "Company": ["TechCorp", "Code Inc."],
        "Application Status": ["Applied", "Interview Scheduled"],
        "Last Update": ["2024-12-01", "2024-12-02"]
    }
    df = pd.DataFrame(data)

    # Save to Excel
    df.to_excel(tracker_file, index=False)
    print(f"Application tracker updated: {tracker_file}")
    return tracker_file

def process_inputs(inputs: dict) -> tuple:
    """
    Process and validate the collected inputs.
    Args:
        inputs (dict): Raw user inputs.
    Returns:
        tuple: (default_testing, job_post_url, resume_doc_filepath) Processed configuration values.
    Raises:
        ValueError: If required inputs are invalid or missing.
    """
    default_testing = inputs.get('default_testing', True)
    job_post_url = inputs.get('job_post_url')
    resume_doc_filepath = inputs.get('resume_doc_filepath')

    if not default_testing:
        # Validate resume file path
        if not resume_doc_filepath:
            raise ValueError("ERROR - Resume file path is required when not in default testing mode.")
        elif not os.path.isfile(resume_doc_filepath):
            raise ValueError(f"ERROR - Resume file '{resume_doc_filepath}' does not exist.")
        
        if not job_post_url:
                    raise ValueError("ERROR - Job post URL is required when not in default testing mode.")
    else:
        resume_doc_filepath = None
        job_post_url = None
        
    return default_testing, job_post_url, resume_doc_filepath

#TODO - make so user jsut spastes it in and explain the program
def integrate_gpt_feedback(resume_text, llm_prompt):
    """
    Simulate GPT output and save the updated resume to a Word document.
    """
    # Simulated GPT response
    gpt_response = f"{resume_text}\n\n(Simulated GPT Edits)\n{llm_prompt}"

    # Save the output
    output_filename = "Updated_Resume.docx"
    save_to_docx(output_filename, gpt_response)
    print(f"Updated resume saved to {output_filename}")
    return output_filename

#TODO
def save_to_docx(filename, content):
    """
    Save text content to a Word document.
    """
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)


def open_file(filepath):
    """WHY - For cross platfrom compatibility"""
    try:
        if sys.platform.startswith('darwin'):
            subprocess.call(('open', filepath))
        elif os.name == 'nt':  # For Windows
            os.startfile(filepath)
        elif os.name == 'posix':  # For Linux
            subprocess.call(('xdg-open', filepath))
    except Exception as e:
        print(f"Unsupported platform: {sys.platform}")
        print(f"Failed to open the file: {e}")
        
def run_program(
    default_testing=True,
    basic_job_description=None,
    basic_resume_text=None,
    resume_doc_filepath=None,
    job_post_url=None
) -> str:
    """
    Run the ATS program to analyze the match between a resume and a job description.

    Args:
    - default_testing (bool): Whether to run in testing mode with basic inputs.
    - basic_job_description (str): Basic job description text for testing.
    - basic_resume_text (str): Basic resume text for testing.
    - resume_doc_filepath (str): File path to a resume in .docx format.
    - job_post_url (str): URL to a job posting.

    Returns:
    - output (str): A formatted string containing the analysis results.
    """
    
    try:
        # Select Inputs
        if default_testing:
            if not basic_job_description or not basic_resume_text:
                raise ValueError("ERROR - Missing basic testing inputs: job description or resume text.")
            job_description_text = basic_job_description
            resume_text = basic_resume_text
        else:
            # Production mode: Load inputs from files and URLs
            if not resume_doc_filepath:
                raise ValueError("ERROR - Resume file path is required in production mode.")
            if not job_post_url:
                raise ValueError("ERROR - Job post URL is required in production mode.")

            resume_text = extract_text_from_docx(resume_doc_filepath)
            job_description_text = extract_text_from_job_description(job_post_url)

            if not resume_text:
                raise ValueError("ERROR - Unable to extract text from resume file.")
            if not job_description_text:
                raise ValueError("ERROR - Unable to extract text from job posting.")

        # Extract keywords
        job_keywords = extract_keywords(job_description_text)
        resume_keywords = extract_keywords(resume_text)

        # Compare keywords
        match_score, matched_keywords, missing_keywords = compare_keywords(job_keywords, resume_keywords)
 
        output = output_results(
            resume_text,
            job_description_text,
            match_score,
            matched_keywords,
            missing_keywords,
            resume_doc_filepath,
            job_post_url
        )
    
        llm_prompt = generate_llm_prompt(resume_text, missing_keywords)
        updated_file = integrate_gpt(resume_text, llm_prompt)
        open_file(updated_file)
        
        return output
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return f"ERROR - Failed to generate output: {str(e)}"
    
#FRONT END
def get_user_input(prompt: str, default_value: any) -> any:
    user_input = input(f"{prompt}: ")
    return user_input.strip() or default_value

def get_user_inputs(resume_doc_filepath: str = None, job_post_url: str = None, default_testing: bool = None) -> dict:
    """
    Collect user inputs for program configuration.
    Returns:
        dict: A dictionary of raw inputs.
    """
    user_input = get_user_input(
        f"Enter Basic Testing Mode (Y/N)? (Default: {'Y' if default_testing else 'N'})",
        default_testing
        )
    
    inputs = {}
    
    # WHY - If the user gives no input the default value is a boolean (or if any non string data) and you can't use .lower() on a boolean so to prevent an error this makes sure it is only done on a string
    if isinstance(user_input, str):
        negatory = ["n", "no", "f", "false", "False"]
        inputs['default_testing'] = user_input.strip().lower() not in negatory
    else:
        inputs['default_testing'] = user_input
        
    if not inputs['default_testing']:
        inputs['resume_doc_filepath'] = get_user_input(
            f"Enter resume file path (Default: {resume_doc_filepath})", 
            resume_doc_filepath
        )
        
        inputs['job_post_url'] = get_user_input(
            f"Enter job posting URL (Default: {job_post_url})", 
            job_post_url
        )
    else:
        inputs['resume_doc_filepath'] = None
        inputs['job_post_url'] = None

    inputs_list = ""
    inputs_list = "\n".join([f"{entry}: {value}" for entry, value in inputs.items()]) 
    logged_inputs = f"""\n
        --- Inputs Collected ---
        {inputs_list}
        -------------------------
        \n"""
    print(logged_inputs)
    
    return inputs

# EXECUTION BLOCK | INTEGRATION - FRONTEND AND BACKEND
if __name__ == "__main__":
    
    #Configure settings / collect user inputs
    raw_inputs = get_user_inputs(resume_doc_filepath, job_post_url, default_testing)
        
    try:
        default_testing, job_post_url, resume_doc_filepath = process_inputs(raw_inputs)
        output = run_program(
            default_testing=default_testing,
            basic_job_description=basic_job_description,
            basic_resume_text=basic_resume_text,
            job_post_url=job_post_url,
            resume_doc_filepath=resume_doc_filepath
        )
        print(output)
    except ValueError as e:
        print(str(e))
